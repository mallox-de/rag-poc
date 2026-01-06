import hashlib
import json
import re
from pathlib import Path
from typing import List, Dict, Any

import time
import uuid
import requests
import snowballstemmer
import typer
from bs4 import BeautifulSoup
from pypdf import PdfReader
from qdrant_client import QdrantClient


from docx import Document
from striprtf.striprtf import rtf_to_text as striprtf_to_text

from qdrant_client.models import (
    VectorParams,
    Distance,
    PointStruct,
    Filter,
    FieldCondition,
    MatchValue,
)
from rank_bm25 import BM25Okapi

app = typer.Typer(add_completion=False)

# ----------------- Defaults (can be overridden by CLI options) -----------------
DEFAULT_QDRANT_URL = "http://localhost:6333"
DEFAULT_COLLECTION = "kb_poc"
DEFAULT_OLLAMA_URL = "http://localhost:11434"
DEFAULT_EMBED_MODEL = "bge-m3"
DEFAULT_CHAT_MODEL = "mistral:7b-instruct"
#DEFAULT_CHAT_MODEL = "qwen2.5:7b-instruct"
#DEFAULT_CHAT_MODEL = "llama3.2"
DEFAULT_DATA_DIR = Path("data")

ALLOWED_SUFFIXES = {".txt", ".md", ".markdown", ".html", ".htm", ".pdf", ".rtf", ".docx"}
IGNORE_DIR_NAMES = {
    ".git", ".svn", ".hg",
    "node_modules", "dist", "build", ".next", ".cache",
    "__pycache__", ".venv", "venv",
}
DEFAULT_MAX_FILE_BYTES = 50 * 1024 * 1024  # 50 MB

DEFAULT_TOP_K_DENSE = 6
DEFAULT_TOP_K_BM25 = 8
DEFAULT_RRF_K = 60

DEFAULT_MAX_CHARS = 900
DEFAULT_OVERLAP = 120

DEFAULT_USE_STEMMING = True

STATE_FILE = Path(".rag_state.json")
BM25_FILE = Path(".bm25_chunks.jsonl")

stemmer_de = snowballstemmer.stemmer("german")
stemmer_en = snowballstemmer.stemmer("english")


# ----------------- Helpers -----------------
def sha1_bytes(b: bytes) -> str:
    return hashlib.sha1(b).hexdigest()


def sha1_text(s: str) -> str:
    return sha1_bytes(s.encode("utf-8", errors="ignore"))


def html_to_text(html: str) -> str:
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def pdf_to_text(path: Path) -> str:
    """
    Extrahiert Text seitenweise aus PDFs.
    Geeignet für Text-PDFs (keine Scans/OCR).
    """
    try:
        reader = PdfReader(str(path))
    except Exception as e:
        typer.secho(f"[WARN] PDF konnte nicht gelesen werden: {path} ({e})", fg=typer.colors.YELLOW)
        return ""

    pages = [page.extract_text() for page in reader.pages]
    pages = [p for p in pages if p and not looks_like_toc_page(p)]

    headers, footers = detect_repeating_lines(pages)

    clean_pages = []
    for p in pages:
        lines = []
        for l in p.splitlines():
            if l.strip() in headers or l.strip() in footers:
                continue
            lines.append(l)
        clean_pages.append("\n".join(lines))

    text = "\n\n".join(clean_pages)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()

def looks_like_toc_page(text: str) -> bool:
    if "inhaltsverzeichnis" not in text.lower():
        return False
    dot_lines = sum(1 for l in text.splitlines() if re.search(r"\.{5,}\s*\d+$", l))
    return dot_lines >= 3

from collections import Counter

def detect_repeating_lines(pages: list[str], top_n=2, threshold=0.3):
    first_lines = Counter()
    last_lines = Counter()

    for p in pages:
        lines = [l.strip() for l in p.splitlines() if l.strip()]
        if not lines:
            continue
        first_lines[lines[0]] += 1
        last_lines[lines[-1]] += 1

    min_count = int(len(pages) * threshold)
    headers = {l for l,c in first_lines.items() if c >= min_count}
    footers = {l for l,c in last_lines.items() if c >= min_count}

    return headers, footers

def looks_like_heading(line: str) -> bool:
    return bool(
        re.match(r"^\d+(\.\d+)*\s+[A-ZÄÖÜ]", line)
        and len(line) < 120
    )

def docx_to_text(path: Path) -> str:
    """
    DOCX -> Text (pure Python via python-docx).
    """
    try:
        doc = Document(str(path))
        parts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        # Optional: Tabellen mitnehmen
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                line = " | ".join([c for c in cells if c])
                if line.strip():
                    parts.append(line)

        text = "\n\n".join(parts)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()
    except Exception as e:
        typer.secho(f"[WARN] DOCX konnte nicht gelesen werden: {path} ({e})", fg=typer.colors.YELLOW)
        return ""


def rtf_to_text(path: Path) -> str:
    """
    RTF -> Text (pure Python via striprtf).
    """
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        text = striprtf_to_text(raw)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()
    except Exception as e:
        typer.secho(f"[WARN] RTF konnte nicht gelesen werden: {path} ({e})", fg=typer.colors.YELLOW)
        return ""


def doc_to_text_best_effort(path: Path) -> str:
    """
    DOC (Word 97-2003) -> Text (Best-effort, pure Python nicht zuverlässig).
    Empfehlung: DOC vorher in DOCX/PDF konvertieren.
    """
    typer.secho(
        f"[WARN] .doc wird im PoC nicht robust (pure Python) unterstützt: {path}. "
        f"Bitte nach .docx oder .pdf konvertieren.",
        fg=typer.colors.YELLOW,
    )
    return ""

def chunk_text(text: str, max_chars: int, overlap: int) -> List[str]:
    """
    PoC: chars statt Token (keine Tokenizer-Abhängigkeit).
    """
    text = text.strip()
    chunks: List[str] = []
    i = 0
    n = len(text)
    while i < n:
        j = min(n, i + max_chars)
        chunk = text[i:j].strip()
        if chunk:
            chunks.append(chunk)
        if j >= n:
            break
        i = max(0, j - overlap)
    return chunks


def tokenize(s: str, use_stemming: bool) -> List[str]:
    # Wörter + Zahlen, inkl. Umlaute (PoC-tauglich)
    toks = re.findall(r"[0-9A-Za-zÄÖÜäöüß]+", s.lower())
    if not use_stemming:
        return toks
    toks = stemmer_de.stemWords(toks)
    toks = stemmer_en.stemWords(toks)
    return toks



# ----------------- Synonyms (Query Expansion) -----------------
def load_synonyms(synonyms_file: Path | None) -> dict[str, list[str]]:
    """
    Lädt eine Synonym-Datei im JSON-Format:
      {
        "sso": ["single sign-on", "einmalanmeldung"],
        "rbac": ["rollenbasierte zugriffskontrolle", "role-based access control"]
      }

    Hinweise:
    - Schlüssel und Synonyme werden case-insensitiv behandelt
    - Datei ist optional (PoC: deterministisch & schnell)
    """
    if not synonyms_file:
        return {}
    try:
        if not synonyms_file.exists():
            typer.secho(f"[WARN] Synonym-Datei nicht gefunden: {synonyms_file}", fg=typer.colors.YELLOW)
            return {}
        data = json.loads(synonyms_file.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            typer.secho(f"[WARN] Synonym-Datei muss ein JSON-Objekt sein: {synonyms_file}", fg=typer.colors.YELLOW)
            return {}

        out: dict[str, list[str]] = {}
        for k, v in data.items():
            if not isinstance(k, str):
                continue
            if isinstance(v, str):
                syns = [v]
            elif isinstance(v, list):
                syns = [x for x in v if isinstance(x, str)]
            else:
                continue
            out[k.strip().lower()] = [s.strip() for s in syns if s and s.strip()]
        return out
    except Exception as e:
        typer.secho(f"[WARN] Synonym-Datei konnte nicht gelesen werden: {synonyms_file} ({e})", fg=typer.colors.YELLOW)
        return {}


def _build_synonym_lookup(syn_map: dict[str, list[str]]) -> dict[str, set[str]]:
    """
    Baut ein Lookup, so dass auch das Auftreten eines Synonyms die gesamte Gruppe erweitert.
    Beispiel:
      "sso" -> {"single sign-on", "einmalanmeldung"}
      "single sign-on" -> {"sso", "einmalanmeldung"}
    """
    lookup: dict[str, set[str]] = {}
    for k, syns in syn_map.items():
        group = {k.lower()} | {s.lower() for s in syns if s}
        for term in group:
            others = group - {term}
            lookup.setdefault(term, set()).update(others)
    return lookup


def expand_query_with_synonyms(question: str, syn_map: dict[str, list[str]]) -> tuple[str, list[str]]:
    """
    Erweitert eine Query deterministisch um bekannte Synonyme.
    Rückgabe: (erweiterte_query, hinzugefügte_terme)

    Matching:
    - single-word Terms: Match via Token-Set
    - multi-word Terms: Match via substring (lowercase)
    """
    if not syn_map:
        return question, []

    q = question.strip()
    q_lower = q.lower()
    tokens = set(re.findall(r"[0-9A-Za-zÄÖÜäöüß]+", q_lower))

    lookup = _build_synonym_lookup(syn_map)

    added: set[str] = set()

    # 1) Matches für single-word tokens
    for tok in list(tokens):
        if tok in lookup:
            added.update(lookup[tok])

    # 2) Matches für multi-word keys/synonyms (substring)
    for term, others in lookup.items():
        if " " in term and term in q_lower:
            added.update(others)

    # Nichts zu tun?
    if not added:
        return question, []

    # Kuratierte Erweiterung ans Ende hängen (leichtgewichtiger als mehrere Queries)
    added_list = sorted({a for a in added if a and a not in q_lower})
    expanded = q + " " + " ".join(added_list)
    return expanded, added_list



def should_skip_path(p: Path, max_file_bytes: int) -> bool:
    if any(part in IGNORE_DIR_NAMES for part in p.parts):
        return True
    if p.suffix.lower() not in ALLOWED_SUFFIXES:
        return True
    try:
        if p.stat().st_size > max_file_bytes:
            print(f"skip because max_file_bytes : {p}")
            return True
    except OSError:
        return True
    return False


def iter_candidate_files(root: Path, max_file_bytes: int):
    for p in root.rglob("*"):
        print(f"Scanning path: {p}")
        if not p.is_file():
            continue
        if should_skip_path(p, max_file_bytes):
            print(f"Skipping path: {p}")
            continue
        yield p


def ollama_embed(ollama_url: str, embed_model: str, text: str) -> List[float]:
    """
    Ollama Embeddings:
    - neuer Endpoint: POST /api/embed
    - Request-Feld: input (string oder array)
    - Response-Feld: embeddings (array von Vektoren)
    """
    r = requests.post(
        f"{ollama_url}/api/embed",
        json={"model": embed_model, "input": text},
        timeout=120,
    )
    r.raise_for_status()
    return r.json()["embeddings"][0]


def ollama_chat(ollama_url: str, chat_model: str, system: str, user: str) -> str:
    payload = {
        "model": chat_model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "stream": False,
    }
    r = requests.post(f"{ollama_url}/api/chat", json=payload, timeout=300)
    r.raise_for_status()
    return r.json()["message"]["content"]


def load_state() -> Dict[str, Any]:
    if STATE_FILE.exists():
        return json.loads(STATE_FILE.read_text(encoding="utf-8"))
    return {"files": {}}


def save_state(state: Dict[str, Any]) -> None:
    STATE_FILE.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")


def ensure_collection(client: QdrantClient, collection: str, dim: int) -> None:
    existing = [c.name for c in client.get_collections().collections]
    if collection in existing:
        return
    client.create_collection(
        collection_name=collection,
        vectors_config=VectorParams(size=dim, distance=Distance.COSINE),
    )


def delete_points_for_source(client: QdrantClient, collection: str, source_path: str) -> None:
    client.delete(
        collection_name=collection,
        points_selector=Filter(
            must=[FieldCondition(key="source_path", match=MatchValue(value=source_path))]
        ),
    )


# ----------------- BM25 persistence -----------------
def bm25_write_rows(rows: List[Dict[str, Any]]) -> None:
    with BM25_FILE.open("a", encoding="utf-8") as f:
        for r in rows:
            f.write(json.dumps(r, ensure_ascii=False) + "\n")


def bm25_rebuild_from_qdrant(client: QdrantClient, collection: str) -> None:
    BM25_FILE.unlink(missing_ok=True)
    scroll_offset = None
    while True:
        points, scroll_offset = client.scroll(
            collection_name=collection,
            with_payload=True,
            with_vectors=False,
            limit=256,
            offset=scroll_offset,
        )
        if not points:
            break
        rows: List[Dict[str, Any]] = []
        for p in points:
            pl = p.payload or {}
            rows.append(
                {
                    "source_path": pl.get("source_path"),
                    "chunk_index": pl.get("chunk_index"),
                    "text": pl.get("text"),
                }
            )
        bm25_write_rows(rows)
        if scroll_offset is None:
            break


def bm25_load_corpus(use_stemming: bool):
    corpus = []
    meta = []
    if not BM25_FILE.exists():
        return corpus, meta
    for line in BM25_FILE.read_text(encoding="utf-8", errors="ignore").splitlines():
        if not line.strip():
            continue
        r = json.loads(line)
        t = r.get("text") or ""
        corpus.append(tokenize(t, use_stemming=use_stemming))
        meta.append(r)
    return corpus, meta


# ----------------- Retrieval -----------------
def dense_search(
    client: QdrantClient,
    collection: str,
    ollama_url: str,
    embed_model: str,
    question: str,
    limit: int,
):
    qvec = ollama_embed(ollama_url, embed_model, question)

    # Qdrant Query API (vereinheitlicht Search/Recommend/Hybrid etc.)
    res = client.query_points(
        collection_name=collection,
        query=qvec,
        limit=limit,
        with_payload=True,
    )

    out = []
    for p in res.points:
        pl = p.payload or {}
        out.append(
            {
                "key": f"{pl.get('source_path')}#{pl.get('chunk_index')}",
                "source_path": pl.get("source_path"),
                "chunk_index": pl.get("chunk_index"),
                "text": pl.get("text", ""),
                "score": float(p.score),
                "origin": "dense",
            }
        )
    return out


_BM25_CACHE = {"bm25": None, "meta": None, "use_stemming": None}

def bm25_search(question: str, limit: int, use_stemming: bool):
    global _BM25_CACHE

    if _BM25_CACHE["bm25"] is None or _BM25_CACHE["use_stemming"] != use_stemming:
        corpus_tok, meta = bm25_load_corpus(use_stemming=use_stemming)
        if not corpus_tok:
            return []
        _BM25_CACHE["bm25"] = BM25Okapi(corpus_tok)
        _BM25_CACHE["meta"] = meta
        _BM25_CACHE["use_stemming"] = use_stemming

    bm25 = _BM25_CACHE["bm25"]
    meta = _BM25_CACHE["meta"]

    qtok = tokenize(question, use_stemming=use_stemming)
    scores = bm25.get_scores(qtok)

    idxs = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:limit]
    out = []
    for i in idxs:
        r = meta[i]
        out.append(
            {
                "key": f"{r.get('source_path')}#{r.get('chunk_index')}",
                "source_path": r.get("source_path"),
                "chunk_index": r.get("chunk_index"),
                "text": r.get("text", ""),
                "score": float(scores[i]),
                "origin": "bm25",
            }
        )
    return out

def rrf_fuse(dense_hits, bm25_hits, rrf_k: int, max_out: int):
    scores: Dict[str, float] = {}
    items: Dict[str, Dict[str, Any]] = {}

    def add_ranked(hits):
        for rank, h in enumerate(hits, start=1):
            key = h["key"]
            scores[key] = scores.get(key, 0.0) + 1.0 / (rrf_k + rank)
            # bevorzugt dense-text falls doppelt, sonst bm25
            if key not in items or items[key].get("origin") != "dense":
                items[key] = h

    add_ranked(dense_hits)
    add_ranked(bm25_hits)

    fused = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)[:max_out]
    out = []
    for key, s in fused:
        h = dict(items[key])
        h["rrf_score"] = s
        out.append(h)
    return out


def build_context_and_refs(fused_hits):
    ctx_parts = []
    refs = []
    for i, h in enumerate(fused_hits, start=1):
        text = (h.get("text", "") or "").strip()

        title_hint = ""
        for line in text.splitlines():
            line = line.strip()
            if len(line) >= 8:
                title_hint = line[:120]
                break

        ctx_parts.append(f"[{i}] {text}")

        refs.append(
            {
                "ref": i,
                "source_path": h.get("source_path"),
                "chunk_index": h.get("chunk_index"),
                "origin": h.get("origin"),
                "rrf_score": h.get("rrf_score"),
                "title_hint": title_hint,
                "text": text,  # <-- HIER ist der Fix
            }
        )
    return ctx_parts, refs


# ----------------- Typer Commands -----------------
@app.command()
def ingest(
    data_dir: Path = typer.Option(DEFAULT_DATA_DIR, "--data-dir", "-d", help="Root-Ordner mit Quellen (rekursiv)."),
    qdrant_url: str = typer.Option(DEFAULT_QDRANT_URL, "--qdrant_storage-url", help="Qdrant URL."),
    collection: str = typer.Option(DEFAULT_COLLECTION, "--collection", help="Qdrant Collection Name."),
    ollama_url: str = typer.Option(DEFAULT_OLLAMA_URL, "--ollama_data-url", help="Ollama URL."),
    embed_model: str = typer.Option(DEFAULT_EMBED_MODEL, "--embed-model", help="Ollama Embedding-Modell."),
    max_chars: int = typer.Option(DEFAULT_MAX_CHARS, "--max-chars", help="Max. Zeichen pro Chunk."),
    overlap: int = typer.Option(DEFAULT_OVERLAP, "--overlap", help="Chunk Overlap (Zeichen)."),
    max_file_bytes: int = typer.Option(DEFAULT_MAX_FILE_BYTES, "--max-file-bytes", help="Max. Dateigröße (Bytes)."),
    rebuild_bm25: bool = typer.Option(True, "--rebuild-bm25/--no-rebuild-bm25", help="BM25 Corpus aus Qdrant rebuilden."),
):
    """
    Ingest:
    - rekursiver Scan (beliebige Tiefe)
    - HTML->Text (Tags entfernen)
    - PDF->Text
    - inkrementelles Reindexing via file_hash
    - Dedupe auf Chunk-Ebene
    - optional BM25-Rebuild für konsistente Hybrid-Suche
    """
    client = QdrantClient(url=qdrant_url)

    # --- Synonym-Expansion (deterministisch, optional) ---
    syn_map = load_synonyms(synonyms_file)
    bm25_question, added_syns = expand_query_with_synonyms(question, syn_map)
    dense_question = bm25_question if expand_dense else question
    if verbose and added_syns:
        typer.secho(f"[INFO] Query expanded with: {', '.join(added_syns)}", fg=typer.colors.CYAN)
        typer.secho(f"[INFO] BM25 query: {bm25_question}", fg=typer.colors.CYAN)
        if expand_dense:
            typer.secho(f"[INFO] Dense query: {dense_question}", fg=typer.colors.CYAN)


    # Dimension einmalig ermitteln
    test_vec = ollama_embed(ollama_url, embed_model, "dimension check")
    ensure_collection(client, collection, dim=len(test_vec))

    state = load_state()
    prev = state.get("files", {})

    scanned = 0
    changed = 0
    points_upserted = 0

    for p in iter_candidate_files(data_dir, max_file_bytes=max_file_bytes):
        scanned += 1
        spath = str(p)
        suffix = p.suffix.lower()

        print( f"p: {p}")

        # load & normalize text based on type
        if suffix in [".html", ".htm"]:
            raw = p.read_text(encoding="utf-8", errors="ignore")
            text = html_to_text(raw)
        elif suffix == ".pdf":
            text = pdf_to_text(p)
        elif suffix == ".docx":
            text = docx_to_text(p)
        elif suffix == ".rtf":
            text = rtf_to_text(p)
        else:
            text = p.read_text(encoding="utf-8", errors="ignore").strip()

        if not text:
            continue

        file_hash = sha1_text(text)

        if prev.get(spath) == file_hash:
            continue  # unchanged

        changed += 1

        # delete all old chunks for file
        delete_points_for_source(client, collection, spath)

        chunks = chunk_text(text, max_chars=max_chars, overlap=overlap)
        seen_chunk_hash = set()

        points = []
        for idx, chunk in enumerate(chunks):
            ch = sha1_text(chunk)
            if ch in seen_chunk_hash:
                continue
            seen_chunk_hash.add(ch)

            vec = ollama_embed(ollama_url, embed_model, chunk)

            # deterministic id per (file + chunkhash)
            #pid = sha1_text(spath + "::" + ch)
            NAMESPACE = uuid.UUID("8d7b2a5f-6f3a-4e3d-9d6a-2b0b2c4e3f10")  # beliebig, aber konstant lassen
            pid = str(uuid.uuid5(NAMESPACE, spath + "::" + ch))

            payload = {
                "source_path": spath,
                "chunk_index": idx,
                "text": chunk,
                "chunk_hash": ch,
                "file_hash": file_hash,
                "file_type": suffix,
            }
            points.append(PointStruct(id=pid, vector=vec, payload=payload))

        if points:
            client.upsert(collection_name=collection, points=points)
            points_upserted += len(points)

        prev[spath] = file_hash

    state["files"] = prev
    save_state(state)

    if rebuild_bm25:
        bm25_rebuild_from_qdrant(client, collection)

    typer.secho(f"Scanned files: {scanned}", fg=typer.colors.CYAN)
    typer.secho(f"Changed files reindexed: {changed}", fg=typer.colors.CYAN)
    typer.secho(f"Upserted chunks: {points_upserted}", fg=typer.colors.CYAN)
    if rebuild_bm25:
        typer.secho("BM25 corpus rebuilt from Qdrant payloads.", fg=typer.colors.GREEN)


@app.command()
def rebuild_bm25(
    qdrant_url: str = typer.Option(DEFAULT_QDRANT_URL, "--qdrant_storage-url", help="Qdrant URL."),
    collection: str = typer.Option(DEFAULT_COLLECTION, "--collection", help="Qdrant Collection Name."),
):
    """Rebuild der lokalen BM25-Corpus-Datei aus Qdrant Payloads."""
    client = QdrantClient(url=qdrant_url)
    bm25_rebuild_from_qdrant(client, collection)
    typer.secho("BM25 corpus rebuilt.", fg=typer.colors.GREEN)


MIN_BM25_SCORE = 1.5   # konservativ, ggf. feinjustieren
MIN_DENSE_SCORE = 0.25

@app.command()
def query(
    question: str = typer.Argument(..., help="Die Frage an die Wissensdatenbank."),
    qdrant_url: str = typer.Option(DEFAULT_QDRANT_URL, "--qdrant_storage-url", help="Qdrant URL."),
    collection: str = typer.Option(DEFAULT_COLLECTION, "--collection", help="Qdrant Collection Name."),
    ollama_url: str = typer.Option(DEFAULT_OLLAMA_URL, "--ollama_data-url", help="Ollama URL."),
    embed_model: str = typer.Option(DEFAULT_EMBED_MODEL, "--embed-model", help="Ollama Embedding-Modell."),
    chat_model: str = typer.Option(DEFAULT_CHAT_MODEL, "--chat-model", help="Ollama Chat-Modell."),
    top_k_dense: int = typer.Option(DEFAULT_TOP_K_DENSE, "--top-k-dense", help="Top-K Dense Treffer."),
    top_k_bm25: int = typer.Option(DEFAULT_TOP_K_BM25, "--top-k-bm25", help="Top-K BM25 Treffer."),
    top_out: int = typer.Option(5, "--top-out", help="Anzahl fusionierter Treffer (Kontext)."),
    rrf_k: int = typer.Option(DEFAULT_RRF_K, "--rrf-k", help="RRF-Konstante."),
    use_stemming: bool = typer.Option(DEFAULT_USE_STEMMING, "--stemming/--no-stemming", help="Stemming für BM25/Tokenisierung."),
    synonyms_file: Path | None = typer.Option(None, "--synonyms-file", help="Optional: Pfad zu synonyms.json für Query-Expansion."),
    expand_dense: bool = typer.Option(False, "--expand-dense", help="Optional: Auch Dense Search mit expandierter Query ausführen."),
    verbose: bool = typer.Option(False, "--verbose", "-v", help="Debug-Ausgaben (z. B. expandierte Query)."),
    show_refs: bool = typer.Option(True, "--refs/--no-refs", help="Referenzen ausgeben."),
):
    """Hybrid Query (Dense + BM25 + RRF) -> LLM Antwort + Referenzen."""
    client = QdrantClient(url=qdrant_url)

    t0 = time.time()
    d = dense_search(client, collection, ollama_url, embed_model, dense_question, top_k_dense)

    t1 = time.time()
    b = bm25_search(bm25_question, top_k_bm25, use_stemming=use_stemming)

    # 1) BM25 prüfen (existiert der Begriff wirklich?)
    max_bm25 = max((h["score"] for h in b), default=0.0)
    max_dense = max((h["score"] for h in d), default=0.0)

    if max_bm25 < MIN_BM25_SCORE and max_dense < MIN_DENSE_SCORE:
        typer.echo(
            "Zu dieser Frage konnten keine relevanten Informationen "
            "in den vorhandenen Quellen gefunden werden."
        )
        return

    query_tokens = set(tokenize(question, use_stemming=True))
    corpus_tokens = set()

    for h in b:
        corpus_tokens.update(tokenize(h["text"], use_stemming=True))

    if not query_tokens & corpus_tokens:
        typer.echo("Der gesuchte Begriff kommt in den Quellen nicht vor.")
        return

    fused = rrf_fuse(d, b, rrf_k=rrf_k, max_out=top_out)

    ctx_parts, refs = build_context_and_refs(fused)

    system = (
        "Du bist ein Assistent für eine lokale Wissensdatenbank.\n"
        "Nutze ausschließlich den bereitgestellten Kontext.\n"
        "Wenn die Antwort nicht im Kontext steht, sage das klar.\n"
        "Wenn mehrere Stellen passen, fasse sie zu einer konsistenten Antwort zusammen.\n"
        "Füge am Ende 'Referenzen:' hinzu und liste jede Referenz als\n"
        "[n] Quelle: <source_path>#<chunk_index> (kurzer Hinweis: <title_hint>)\n"
        "Gib KEINE Referenzen an, die nicht im Kontext vorkommen."
    )

    user = (
        f"FRAGE:\n{question}\n\n"
        "Fasse den Text präzise zusammen. Keine Bewertung, keine Warnhinweise, keine Ablehnung."
        "Nutze nur Informationen aus dem Kontext.\n\n"
        f"KONTEXT:\n" + "\n\n".join(ctx_parts)
    )

    t2 = time.time()
    answer = ollama_chat(ollama_url, chat_model, system, user)

    t3 = time.time()
    print(f"dense={t1 - t0:.2f}s bm25={t2 - t1:.2f}s llm={t3 - t2:.2f}s total={t3 - t0:.2f}s")

    # 1) Antwort ausgeben
    typer.echo(answer)

    # 2) Referenzen IMMER deterministisch ausgeben (nicht dem LLM überlassen)
    if show_refs:
        MAX_PRINT_CHARS = 300

        typer.echo("\n\nReferenzen:")
        for r in refs:
            src = r.get("source_path", "?")
            idx = r.get("chunk_index", "?")
            origin = r.get("origin", "?")
            score = r.get("rrf_score", None)
            text = r.get("text", "")
            hint = r.get("title_hint", "")

            score_txt = f"{score:.6f}" if isinstance(score, (int, float)) else "?"
            typer.echo(f"\n[{r['ref']}] {src}#{idx} ({origin}, rrf={score_txt}) - {hint}")
            typer.echo("-" * 80)


            text = text[:MAX_PRINT_CHARS] + "…" if len(text) > MAX_PRINT_CHARS else text
            typer.echo(text)

        # 2) Referenzen + Chunk-Text direkt aus fused (Quelle!)
        typer.echo("\n\nReferenzen (mit Chunk-Text):")
        for i, h in enumerate(fused, start=1):
            src = h.get("source_path", "?")
            idx = h.get("chunk_index", "?")
            origin = h.get("origin", "?")
            rrf = h.get("rrf_score", None)
            text = (h.get("text", "") or "").strip()

            rrf_txt = f"{rrf:.6f}" if isinstance(rrf, (int, float)) else "?"
            typer.echo(f"\n[{i}] {src}#{idx} ({origin}, rrf={rrf_txt})")
            typer.echo("-" * 80)
            text = text[:MAX_PRINT_CHARS] + "…" if len(text) > MAX_PRINT_CHARS else text
            typer.echo(text)

@app.command()
def reset_state(
    delete_bm25: bool = typer.Option(False, "--delete-bm25", help="BM25-Datei (.bm25_chunks.jsonl) ebenfalls löschen."),
):
    """Löscht den lokalen State (Hashes). Nützlich für Full Reindex."""
    if STATE_FILE.exists():
        STATE_FILE.unlink()
        typer.secho("Deleted .rag_state.json", fg=typer.colors.GREEN)
    else:
        typer.secho(".rag_state.json not found", fg=typer.colors.YELLOW)

    if delete_bm25:
        if BM25_FILE.exists():
            BM25_FILE.unlink()
            typer.secho("Deleted .bm25_chunks.jsonl", fg=typer.colors.GREEN)
        else:
            typer.secho(".bm25_chunks.jsonl not found", fg=typer.colors.YELLOW)


if __name__ == "__main__":
    app()
